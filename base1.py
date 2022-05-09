# в файле с названием Фамилии содержится список фамилий – каждая фамилия через enter
#создается файл uuu где на каждом листе написана фамилия и вариант
from docx import Document
from docx.shared import Inches
import random as rand
from sympy import Rational
a=0
document = Document("Фамилии.docx")
docum = Document()
for paragraph in document.paragraphs:
    a+=1
    docum.add_paragraph("Вариант № \u222B\u2080\u221E x^2=x\u00B2 \u1D44"+str(a)+" - "+paragraph.text)
    docum.add_page_break()    
docum.save("uuu.docx")

answer = [0 for i in range(7)]#список для ответов
def task1(var):
    if var%2==1:
        isTask = rand.randint(1,2)#какое из двух заданий
        isTask=2
        isPoint = rand.randint(1,2)#пункт а или б
        isPoint=2
        taskFirst = ["""Имеется пять отрезков, длины которых соответственно равны""",
                     """ Наугад берут три из них. Какова вероятность того, что """]
        taskFirstA = ["первый отрезок будет длиной ",
                      ", а второй — "]
        taskFirstB = "из этих отрезков можно построить треугольник?"
        taskSecond = ["Среди десяти подарков к Новому году ","""с красной икрой,""",
                      " — с черной и ",
                      """ с икрой заморской, баклажанной. Какова вероятность того, что среди трех наугад взятых подарков """]
        taskSecondA = "два содержат красную икру?"
        taskSecondB = "все три подарка с разной икрой?"
        
        if isTask==1:#если выпала первая задача
            cifr11=[]
            for i in range(5):
                r=rand.randint(1,15)
                while r in cifr11:
                    r=rand.randint(1,15)
                cifr11.append(r)
            s=""
            if isPoint==1:#рассматриваем разные пункты
                s+=taskFirstA[0]+str(cifr11[0])+taskFirstA[1]+str(cifr11[1])+" см."
            else:
                s=taskFirstB

            cifr11.sort()
            
            if isPoint==1:#ответы
                answer[0] = "1/20"
            else:
                kl=0
                print(cifr11)
                for i in range(3):
                    for j in range(i+1,4):
                        for k in range(j+1,5):
                            if cifr11[i]+cifr11[j]>cifr11[k]:
                                if cifr11[i]+cifr11[k]>cifr11[j]:
                                    if cifr11[j]+cifr11[k]>cifr11[i]:
                                        kl+=1
                answer[0]=str(kl/10)
                            
                
            listCifr = "" #строка с цифрами
            for i in range(4):
                listCifr+=" "+str(cifr11[i])+","
            listCifr += " "+str(cifr11[4])+' см.'

            ex = taskFirst[0]+listCifr+taskFirst[1]+s
            print(ex)
            
        elif isTask==2:
            cifr11=[2,2,2]#закинули в минимум по 2 подарка
            r=rand.randint(0,4)
            cifr11[0]+=r#сделали разные значения для разного вида икры
            cifr11[1]+=4-r
            s=""#разбиваем на пункты
            if isPoint==1:
                result = 1
                result*=(cifr11[0]/10)*((cifr11[0]-1)/9)*((cifr11[1]+cifr11[2])/8)
                answer[0]=str(round(result,3))
                s+=taskSecondA
            else:
                s=taskSecondB
                result=1
                result*=(cifr11[0]/10)*(cifr11[1]/9)*(cifr11[2]/8)
                answer[0]=str(round(result,3))
            ex = taskSecond[0]+str(cifr11[0])
            if cifr11[0]>=5:
                ex+=" подарков "
            else:
                ex+=" подарка "
            ex+=taskSecond[1]+str(cifr11[1])+taskSecond[2]+str(cifr11[2])+taskSecond[3]
            ex+=s
            print(ex)

            
                
                
                
            
            
            
            
