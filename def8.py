from sympy import sin,pi
from sympy.abc import x
import random as rand
from docx import Document
docum=Document()

def task8(var):
    docum.add_picture("picture\def8_0.png")
    if var%2==1:
        docum.add_picture("picture\def8_1.png")
        ans = "\t       0, x\u22642;\nf(x) = \t2x/5, 2<x\u22643;\n\t       0, x>3;"
        ans+="\nP(2<X<2,5) = F(2,5)-F(2)=0,45"
        docum.add_paragraph(ans)
        
    else:
        docum.add_picture("picture\def8_2.png")
        ans = "\t        0, x\u22640;\nf(x) =\tcos(x), 0<x\u2264\u03C0/2;"
        ans+="\n\t        0, x>\u1D745/2"
        ans+="\nP(0<x<\\u03C0/6) = F(\u03C0/6)-F(0) = 1/2"
        docum.add_paragraph(ans)
    
    

task8(0)
task8(1)
docum.save("uu2.docx")
