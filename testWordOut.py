
from docx import Document
from docx.shared import Inches

document = Document()
document.add_heading('Document Title', 0)
p = document.add_paragraph('A plain paragraph having some ')   # Insert the paragraph and define the paragraph as a P variable
p.add_run('bold').bold = True   # Add bold word bold
p.add_run(' and some ')         # Add ordinary word
p.add_run('italic.').italic = True   #

document.add_heading('Heading, level 1', level=1)          # 1 grade title, 1 point than 0
document.add_paragraph('Intense quote', style='IntenseQuote')   # Add a paragraph, style is emphasized
document.add_page_break()
document.add_paragraph(
    'first item in unordered list', style='ListBullet'          # Add a paragraph, have a point in front
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'            # Add a paragraph, there is a serial number 1
)

#document.add_picture('monty-truth.png', width=Inches(1.25))      # Insert the picture

table = document.add_table(rows=1, cols=3)                      # Insert the table
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
recordset =[{'qty':'1a','id':1,'desc':'aaa'},{'qty':'2a','id':2,'desc':'bbb'},{'qty':'3c','id':3,'desc':'ccc'}]
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item['qty'])
    row_cells[1].text = str(item['id'])
    row_cells[2].text = item['desc']

document.add_page_break()    # Add Next       

document.save('demo.docx')  # Save the document

