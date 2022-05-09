from docx import Document
docum=Document()

def task9(var):
    docum.add_picture("picture\def9_0.png")
    if var%2==1:
        docum.add_picture("picture\def9_1.png")
        ans ="a=0,5"
        ans+="\n\t         0, x<2;\nF(x) =\tx\u00B2-x-2, 2\u2264x\u22643;\n\t         1, x>3."
        docum.add_paragraph(ans)
        
    else:
        docum.add_picture("picture\def8_2.png")
        ans = "a=1"
        ans+="\n\t         0, x<0;\nF(x) =\t(-cos(x)+1)/2, 0\u2264x\u2264\u03C0/2;\n\t         1, x>\u03C0/2."

        docum.add_paragraph(ans)
        
task9(0)
task9(1)
docum.save("uu2.docx")
