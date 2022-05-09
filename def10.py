from docx import Document
docum = Document()

def task10(var):
    docum.add_picture("picture\def10_0.png")
    if var%2==1:
        docum.add_picture("picture\def10_1.png")
        ans = "\n\t                       0, x\u22643;\nF(x) =\t1/8(x\u00B2-6x+9), 3<x\u22645"
        ans+="\n\t             x/2 - 2, 5<x\u22646;\n\t        1, x>6."
        ans+="\nP(4\u2264x\u22645,5)=5/8=0,625"
        ans+="\nM(x) = 59/12=4,92\nD(x) = 74/3 - (59/12)^2=0,493\n\u03C3(x)=0,702"
        docum.add_paragraph(ans)
    else:
        docum.add_picture("picture\def10_2.png")
        ans="\n\t\t         0, x\u22641;\nF(x)   = \t1/3(x\u00B2-2x+1), 1<x\u22642,5;"
        ans+="\n\t          -x\u00B2+6x-8, 2,5<x\u22643;\n\t\t         1,x>3"
        ans+="\nP(0\u2264x\u22642,5)=3/4"
        ans+="\nM(x) = 13/6=2,167\nD(x) = 39/8 - (13/6)^2=13/72=0,181\n\u03C3(x)=0,425"
        docum.add_paragraph(ans)
task10(0)
task10(1)
docum.save("uu2.docx")       
